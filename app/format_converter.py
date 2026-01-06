"""
Format Converter - Convert between PDF, DOCX, and TXT formats
Stable, page-aware implementation
"""

# =======================
# IMPORTS
# =======================

import re
from datetime import datetime

import PyPDF2
  # PyMuPDF

from docx import Document
from docx.shared import Pt

from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_CENTER

from xml.sax.saxutils import escape


# =======================
# UTILITIES
# =======================

def clean_xml_text(text: str) -> str:
    """
    Remove NULL bytes and illegal XML characters
    """
    if not text:
        return ""

    text = text.replace("\x00", "")
    text = re.sub(r'[\x00-\x08\x0B-\x0C\x0E-\x1F\x7F]', '', text)

    try:
        text = text.encode("utf-8", errors="ignore").decode("utf-8")
    except Exception:
        pass

    return text


# =======================
# PDF TEXT EXTRACTION
# =======================

def pdf_to_text(pdf_path: str) -> str:
    """
    Convert PDF to plain text (legacy / fallback)
    """
    try:
        output = []
        with open(pdf_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                output.append(page.extract_text() or "")
        return "\n".join(output)
    except Exception as e:
        raise Exception(f"Failed to convert PDF to text: {e}")


def pdf_to_text_by_page(pdf_path: str) -> dict:
    """
    Extract PDF text page-by-page
    Returns: { page_index: text }
    """
    pages = {}
    try:
        with open(pdf_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for i, page in enumerate(reader.pages):
                pages[i] = clean_xml_text(page.extract_text() or "")
        return pages
    except Exception as e:
        raise Exception(f"Failed to extract PDF text by page: {e}")


# =======================
# PDF IMAGE EXTRACTION
# =======================



# =======================
# PDF → DOCX
# =======================

def pdf_to_docx(pdf_path: str, output_path: str) -> str:
    """
    Convert PDF to DOCX (text-only, page-aware, stable)
    """
    try:
        page_text = pdf_to_text_by_page(pdf_path)

        if not any(page_text.values()):
            doc = Document()
            doc.add_heading("Scanned PDF Detected", level=0)
            doc.add_paragraph(
                "This PDF contains no machine-readable text. "
                "It appears to be a scanned document."
            )
            doc.save(output_path)
            return output_path

        doc = Document()
        doc.add_heading("Converted Document", level=0)
        doc.add_paragraph(
            f"Converted from PDF on {datetime.now().strftime('%B %d, %Y')}"
        )

        for page_index in page_text:
            paragraphs = page_text[page_index].split("\n")

            for para in paragraphs:
                para = clean_xml_text(para.strip())
                if not para:
                    continue

                # Heading heuristic
                if para.isupper() and len(para.split()) <= 8:
                    doc.add_heading(para, level=1)
                    continue

                # Bullet heuristic
                if para.startswith(("-", "•", "*")):
                    doc.add_paragraph(
                        para.lstrip("-•* ").strip(),
                        style="List Bullet"
                    )
                    continue

                p = doc.add_paragraph(para)
                p.paragraph_format.space_after = Pt(8)

        doc.save(output_path)
        return output_path

    except Exception as e:
        raise Exception(f"Failed to convert PDF to DOCX: {e}")


# =======================
# DOCX ↔ TEXT
# =======================

def docx_to_text(docx_path: str) -> str:
    try:
        doc = Document(docx_path)
        return "\n".join(p.text for p in doc.paragraphs)
    except Exception as e:
        raise Exception(f"Failed to convert DOCX to text: {e}")


def text_to_docx(text: str, output_path: str, title: str = "Document") -> str:
    try:
        doc = Document()
        doc.add_heading(title, level=0)
        doc.add_paragraph(
            f"Created on {datetime.now().strftime('%B %d, %Y %I:%M %p')}"
        )

        for line in text.split("\n"):
            doc.add_paragraph(clean_xml_text(line))

        doc.save(output_path)
        return output_path
    except Exception as e:
        raise Exception(f"Failed to convert text to DOCX: {e}")


# =======================
# TEXT → PDF
# =======================

def text_to_pdf(text: str, output_path: str, title: str = "Document") -> str:
    try:
        doc = SimpleDocTemplate(output_path, pagesize=letter)
        styles = getSampleStyleSheet()
        elements = []

        title_style = ParagraphStyle(
            "Title",
            parent=styles["Heading1"],
            alignment=TA_CENTER
        )

        elements.append(Paragraph(escape(title), title_style))
        elements.append(Spacer(1, 12))

        for i, line in enumerate(text.split("\n")):
            if line.strip():
                elements.append(Paragraph(escape(line), styles["Normal"]))
                elements.append(Spacer(1, 6))

            if i > 0 and i % 50 == 0:
                elements.append(PageBreak())

        doc.build(elements)
        return output_path

    except Exception as e:
        raise Exception(f"Failed to convert text to PDF: {e}")


# =======================
# MAIN CONVERTER (API SAFE)
# =======================

def convert_format(input_path: str, output_path: str,
                   from_format: str, to_format: str,
                   title: str = "Document") -> str:
    """
    Central conversion dispatcher (DO NOT CHANGE SIGNATURE)
    """
    from_format = from_format.lower()
    to_format = to_format.lower()

    if from_format == to_format:
        raise Exception("Source and target formats are the same")

    if from_format == "pdf":
        if to_format == "txt":
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(pdf_to_text(input_path))
            return output_path

        if to_format == "docx":
            return pdf_to_docx(input_path, output_path)

    if from_format == "docx":
        if to_format == "txt":
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(docx_to_text(input_path))
            return output_path

        if to_format == "pdf":
            return text_to_pdf(docx_to_text(input_path), output_path, title)

    if from_format == "txt":
        with open(input_path, "r", encoding="utf-8") as f:
            text = f.read()

        if to_format == "docx":
            return text_to_docx(text, output_path, title)

        if to_format == "pdf":
            return text_to_pdf(text, output_path, title)

    raise Exception(f"Unsupported format conversion: {from_format} → {to_format}")
